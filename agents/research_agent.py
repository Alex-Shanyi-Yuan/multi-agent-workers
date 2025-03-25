from typing import Dict, Any, List, Optional
import os
from pathlib import Path

from autogen_ext.models.openai import OpenAIChatCompletionClient

from .base_agent import BaseAgent

import PyPDF2
import docx
from datetime import datetime
import logging

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class ResearchAgent(BaseAgent):
    """Agent responsible for searching through documents and finding relevant information."""
    
    def __init__(self,
        model_client: Optional[OpenAIChatCompletionClient],
        handoffs: List[str],
    ):
        system_message = """
            You are the **Knowledge Retrieval Agent**. Your role is to fetch data from Confluence, PDFs, or Word docs.

            **Policies**:
            1. Never invent information - respond "No relevant documents found" if unsure.

            **Steps**:
            1. Receive the query.
            2. Identify the files that may contain information using `get_avalible_files`.
            3. Identify type of document (PDF, Word).
            4. Get the content by calling `extract_text`.
            5. Summarize findings and cite sources.
            6. Use TERMINATE when reaseaching is complete.

            **Notes**:
            - Use `request_document_access` if files are restricted.
            - Flag outdated documents with `flag_obsolete_content`.
        """
        
        super().__init__(
            name="Researcher",
            system_message=system_message,
            model_client=model_client,
            handoffs=handoffs,
            tools=[
                self.extract_pdf_text,
                self.extract_docx_text,
                self.get_available_files,
                # self.confluence.search,
            ]
        )

    def get_available_files(self) -> List[str]:
        """Get a list of available files."""

        assets_folder = os.path.join(os.path.dirname(__file__), '../', 'assets')
        files = [os.path.join(assets_folder, file) for file in os.listdir(assets_folder)]
        return files

    def _get_pdf_content(
        self,
        pdf_path: Path,
    ) -> Optional[Dict[str, Any]]:
        """Get PDF content by extraction.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Dict containing PDF content and metadata
        """
        try:
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                content = {
                    "pages": [],
                    "metadata": reader.metadata,
                    "page_count": len(reader.pages),
                    "title": reader.metadata.get('/Title', pdf_path.name),
                    "extracted_at": datetime.now().isoformat()
                }
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text()
                        content["pages"].append({
                            "page_number": page_num + 1,
                            "text": text
                        })
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        
                return content
                
        except Exception as e:
            logger.error(f"Error extracting content from {pdf_path}: {str(e)}")
            return None
    
    def extract_pdf_text(
        self,
        pdf_path: str,
    ) -> Optional[Dict[str, Any]]:
        """Extract all text from a PDF file.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Dict containing extracted text and metadata
        """
        return self._get_pdf_content(Path(pdf_path))
    
    def _get_docx_content(self, docx_path: Path) -> Optional[Dict[str, Any]]:
        try:
            doc = docx.Document(docx_path)
            
            content = {
                "paragraphs": [],
                "metadata": {
                    "core_properties": {
                        prop: str(value)
                        for prop, value in doc.core_properties.__dict__.items()
                        if not prop.startswith('_')
                    }
                },
                "paragraph_count": len(doc.paragraphs),
                "title": doc.core_properties.title or docx_path.name,
                "extracted_at": datetime.now().isoformat()
            }
            
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    content["paragraphs"].append({
                        "index": i,
                        "text": para.text,
                        "style": para.style.name,
                        "level": para.style.base_style.name if para.style.base_style else None
                    })
                        
            return content
                
        except Exception as e:
            logger.error(f"Error extracting content from {docx_path}: {str(e)}")
            return None

    def extract_docx_text(
        self,
        docx_path: str,
    ) -> Optional[Dict[str, Any]]:
        """Extract all text from a Word document.
        
        Args:
            docx_path: Path to Word document
            
        Returns:
            Dict containing extracted text and metadata
        """
        return self._get_docx_content(Path(docx_path))
