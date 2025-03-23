import os
import logging
from typing import Dict, List, Any, Optional, Generator
from pathlib import Path
import docx
import re
from datetime import datetime
import json
from tqdm import tqdm

logger = logging.getLogger(__name__)

class DocxProcessor:
    """Processor for searching and extracting content from Word documents."""
    
    def __init__(self, cache_dir: Optional[str] = None):
        """Initialize the Word document processor.
        
        Args:
            cache_dir: Directory for caching extracted document content
        """
        self.cache_dir = cache_dir or os.getenv('DOCUMENT_CACHE_DIR', './cache/docx')
        Path(self.cache_dir).mkdir(parents=True, exist_ok=True)
        
        self.max_file_size = int(os.getenv('MAX_DOCX_SIZE_MB', '10')) * 1024 * 1024  # Convert to bytes
        
    def search(
        self,
        query: str,
        docx_dir: str,
        recursive: bool = True,
        use_cache: bool = True
    ) -> List[Dict[str, Any]]:
        """Search for text in Word documents.
        
        Args:
            query: Search query
            docx_dir: Directory containing Word documents
            recursive: Whether to search subdirectories
            use_cache: Whether to use cached content
            
        Returns:
            List of search results with metadata
        """
        results = []
        query_pattern = re.compile(query, re.IGNORECASE)
        
        # Find all Word documents
        docx_files = self._find_docx_files(docx_dir, recursive)
        
        for docx_file in tqdm(docx_files, desc="Searching Word documents"):
            try:
                # Check file size
                if os.path.getsize(docx_file) > self.max_file_size:
                    logger.warning(f"Skipping {docx_file}: File too large")
                    continue
                    
                # Get content (from cache or by extraction)
                content = self._get_docx_content(docx_file, use_cache)
                if not content:
                    continue
                    
                # Search for matches
                matches = self._find_matches(content, query_pattern)
                if matches:
                    results.append({
                        "file_path": str(docx_file),
                        "title": content.get("title", os.path.basename(docx_file)),
                        "matches": matches,
                        "metadata": content.get("metadata", {}),
                        "paragraph_count": content.get("paragraph_count", 0)
                    })
                    
            except Exception as e:
                logger.error(f"Error processing {docx_file}: {str(e)}")
                
        return results
        
    def _find_docx_files(self, directory: str, recursive: bool) -> Generator[Path, None, None]:
        """Find Word documents in directory.
        
        Args:
            directory: Directory to search
            recursive: Whether to search subdirectories
            
        Yields:
            Path objects for Word documents
        """
        directory = Path(directory)
        pattern = '**/*.docx' if recursive else '*.docx'
        
        for docx_file in directory.glob(pattern):
            if docx_file.is_file():
                yield docx_file
                
    def _get_docx_content(
        self,
        docx_path: Path,
        use_cache: bool
    ) -> Optional[Dict[str, Any]]:
        """Get Word document content, either from cache or by extraction.
        
        Args:
            docx_path: Path to Word document
            use_cache: Whether to use cached content
            
        Returns:
            Dict containing document content and metadata
        """
        cache_file = Path(self.cache_dir) / f"{docx_path.stem}_{docx_path.stat().st_mtime}.json"
        
        # Check cache first
        if use_cache and cache_file.exists():
            try:
                with cache_file.open('r') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"Error reading cache for {docx_path}: {str(e)}")
                
        try:
            doc = docx.Document(docx_path)
            
            # Extract text and metadata
            content = {
                "paragraphs": [],
                "metadata": {
                    "core_properties": {
                        prop: str(value)
                        for prop, value in doc.core_properties.__dict__.items()
                        if not prop.startswith('_')
                    }
                },
                "paragraph_count": len(doc.paragraphs),
                "title": doc.core_properties.title or docx_path.name,
                "extracted_at": datetime.now().isoformat()
            }
            
            # Extract text from each paragraph with style information
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():  # Skip empty paragraphs
                    content["paragraphs"].append({
                        "index": i,
                        "text": para.text,
                        "style": para.style.name,
                        "level": para.style.base_style.name if para.style.base_style else None
                    })
                    
            # Cache the content
            with cache_file.open('w') as f:
                json.dump(content, f)
                
            return content
            
        except Exception as e:
            logger.error(f"Error extracting content from {docx_path}: {str(e)}")
            return None
            
    def _find_matches(
        self,
        content: Dict[str, Any],
        pattern: re.Pattern
    ) -> List[Dict[str, Any]]:
        """Find pattern matches in Word document content.
        
        Args:
            content: Extracted document content
            pattern: Compiled regex pattern
            
        Returns:
            List of matches with context
        """
        matches = []
        
        for para in content["paragraphs"]:
            para_text = para["text"]
            para_matches = pattern.finditer(para_text)
            
            for match in para_matches:
                # Get context around match
                start = max(0, match.start() - 100)
                end = min(len(para_text), match.end() + 100)
                
                matches.append({
                    "paragraph_index": para["index"],
                    "style": para["style"],
                    "match": match.group(),
                    "context": para_text[start:end],
                    "position": {
                        "start": match.start(),
                        "end": match.end()
                    }
                })
                
        return matches
        
    def extract_text(
        self,
        docx_path: str,
        use_cache: bool = True
    ) -> Optional[Dict[str, Any]]:
        """Extract all text from a Word document.
        
        Args:
            docx_path: Path to Word document
            use_cache: Whether to use cached content
            
        Returns:
            Dict containing extracted text and metadata
        """
        return self._get_docx_content(Path(docx_path), use_cache)
        
    def clear_cache(self, older_than_days: Optional[int] = None):
        """Clear the Word document content cache.
        
        Args:
            older_than_days: Only clear cache older than this many days
        """
        try:
            cache_path = Path(self.cache_dir)
            for cache_file in cache_path.glob('*.json'):
                if older_than_days:
                    # Check file age
                    file_age = datetime.now() - datetime.fromtimestamp(cache_file.stat().st_mtime)
                    if file_age.days > older_than_days:
                        cache_file.unlink()
                else:
                    cache_file.unlink()
        except Exception as e:
            logger.error(f"Error clearing cache: {str(e)}")
            
    def get_metadata(self, docx_path: str) -> Optional[Dict[str, Any]]:
        """Get metadata from a Word document.
        
        Args:
            docx_path: Path to Word document
            
        Returns:
            Dict containing document metadata
        """
        try:
            doc = docx.Document(docx_path)
            return {
                "core_properties": {
                    prop: str(value)
                    for prop, value in doc.core_properties.__dict__.items()
                    if not prop.startswith('_')
                },
                "paragraph_count": len(doc.paragraphs),
                "file_size": os.path.getsize(docx_path),
                "last_modified": datetime.fromtimestamp(
                    os.path.getmtime(docx_path)
                ).isoformat()
            }
        except Exception as e:
            logger.error(f"Error getting metadata from {docx_path}: {str(e)}")
            return None
